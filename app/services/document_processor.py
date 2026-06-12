import logging
import os
import zipfile
import xml.etree.ElementTree as ET
import docx
from PyPDF2 import PdfReader

from sqlalchemy.orm import Session
from app.core.config import settings
from app.models.document import Document
from app.services.vector_store import VectorStore

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: str) -> tuple[list[dict], int]:
   
    pages_data = []
    
    # 1. Check file path valid
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found at path: {file_path}")
        
    # 2. PyPDF2 PdfReader se PDF read
    reader = PdfReader(file_path)
    total_pages = len(reader.pages)
    
    # 3. every page text read 
    for page_num, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            pages_data.append({
                "text": text,
                "page_number": page_num + 1
            })
            
    return pages_data, total_pages

def extract_text_from_txt(file_path: str) -> tuple[list[dict], int]:
    """
    Extract text from a plain TXT file.
    Creates virtual pages of 2000 characters each.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"TXT file not found at path: {file_path}")
    
    # Encodings handle karo (UTF-8, UTF-8-sig, or CP1252 as fallback)
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, "r", encoding="utf-8-sig") as f:
                full_text = f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="cp1252") as f:
                full_text = f.read()
    pages_data = []
    char_per_page = 2000
    text_len = len(full_text)
    
    if text_len == 0:
        return [], 0
    start = 0
    page_num = 1
    while start < text_len:
        end = start + char_per_page
        page_text = full_text[start:end]
        pages_data.append({
            "text": page_text,
            "page_number": page_num
        })
        start = end
        page_num += 1
        
    return pages_data, page_num - 1
def extract_text_from_docx(file_path: str) -> tuple[list[dict], int]:
    """
    Extract text from a Word DOCX file using python-docx.
    If python-docx fails (e.g. due to corrupted media files / bad CRC-32),
    falls back to direct XML extraction from the ZIP container.
    Creates virtual pages of 2000 characters each.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found at path: {file_path}")
        
    paragraphs_text = []

    try:
        doc = docx.Document(file_path)
        # Read Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs_text.append(para.text.strip())
                
        # Read Tables (if any)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs_text.append(" | ".join(row_text))
    except Exception as e:
        logger.warning(f"python-docx failed to read docx, attempting direct XML extraction fallback: {e}")
        try:
            with zipfile.ZipFile(file_path) as z:
                # Read main document XML directly
                doc_xml = z.read('word/document.xml')
                
            root = ET.fromstring(doc_xml)
            namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
            
            # Extract paragraphs
            for p in root.findall('.//w:p', namespaces):
                p_text = "".join([t.text for t in p.findall('.//w:t', namespaces) if t.text])
                if p_text.strip():
                    paragraphs_text.append(p_text.strip())
        except Exception as fallback_err:
            logger.error(f"Fallback direct XML docx parsing failed: {fallback_err}")
            raise e

    full_text = "\n\n".join(paragraphs_text)
    
    pages_data = []
    char_per_page = 2000
    text_len = len(full_text)
    
    if text_len == 0:
        return [], 0
    start = 0
    page_num = 1
    while start < text_len:
        end = start + char_per_page
        page_text = full_text[start:end]
        pages_data.append({
            "text": page_text,
            "page_number": page_num
        })
        start = end
        page_num += 1
        
    return pages_data, page_num - 1

def chunk_text(pages_data: list[dict], chunk_size: int, chunk_overlap: int) -> tuple[list[str], list[dict]]:
   
    chunks = []
    metadatas = []
    
    # Har page ka text process karo
    for page in pages_data:
        text = page["text"]
        page_num = page["page_number"]
        
        # Sliding Window algorithm character for overlapping
        start = 0
        text_len = len(text)
        
        while start < text_len:
            # End index target block size limits par set
            end = start + chunk_size
            chunk_content = text[start:end]
            
            chunks.append(chunk_content)
            metadatas.append({
                "page_number": page_num,
                "length": len(chunk_content)
            })
            
            # Start position shift , overlap buffer subtract karke
            start += chunk_size - chunk_overlap
            
    return chunks, metadatas


def process_document(db: Session, document_id: int) -> None:
  
    # 1. Database se core document details fetch 
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        logger.error(f"Document ID {document_id} not found in database.")
        return

    try:
        # 2. Document state 'processing' mark 
        document.status = "processing"
        db.commit()
        logger.info(f"Started processing document ID {document_id} ({document.original_filename}).")

        # 3. File type check karke appropriate text extract karo
        _, ext = os.path.splitext(document.storage_path.lower())
        if ext == ".pdf":
            pages_data, total_pages = extract_text_from_pdf(document.storage_path)
        elif ext == ".txt":
            pages_data, total_pages = extract_text_from_txt(document.storage_path)
        elif ext == ".docx":
            pages_data, total_pages = extract_text_from_docx(document.storage_path)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
        
        # 4. Text dynamic chunks splits
        chunks, metadatas = chunk_text(
            pages_data=pages_data,
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP
        )

        # 5. Vector Store (ChromaDB) mein add/index karo
        vs = VectorStore()
        vs.add_document_chunks(
            document_id=document.id,
            chunks=chunks,
            metadatas=metadatas
        )

        # 6. Database record successfully complete mark
        document.status = "completed"
        document.page_count = total_pages
        document.chunk_count = len(chunks)
        document.error_message = None
        db.commit()
        
        logger.info(f"Successfully processed document ID {document_id}. Pages: {total_pages}, Chunks: {len(chunks)}.")

    except Exception as e:
        # 7. DB clean-up and fail state set after getting any error
        logger.exception(f"Error processing document ID {document_id}: {str(e)}")
        db.rollback()
        
        document.status = "failed"
        document.error_message = str(e)
        db.commit()
